"""Body content builders: chapters, sections, tables, bibliography, appendices."""
import base64
import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

from schemas.thesis_schema import Appendix, Chapter, FigureEntry, Section, TableEntry
from services.docx_primitives import (
    BODY_PT, _add_page_break, _body, _center, _chapter_heading,
    _heading1, _heading2, _para_fmt, _set_font, _to_roman, _to_zh,
)


# ── Figure rendering ──────────────────────────────────────────────────────────

def render_figure(doc: Document, figures: list[FigureEntry], marker_idx: int) -> None:
    if marker_idx >= len(figures):
        return
    entry = figures[marker_idx]
    if not entry.images:
        return

    for img in entry.images:
        img_stream = io.BytesIO(base64.b64decode(img.imageData))
        para = doc.add_paragraph()
        _para_fmt(para, align=WD_ALIGN_PARAGRAPH.CENTER, space_before_pt=4)
        para.add_run().add_picture(img_stream, width=Inches(5))

    caption = doc.add_paragraph(style="FigCaption")
    _para_fmt(caption, align=WD_ALIGN_PARAGRAPH.CENTER, space_after_pt=6)
    _set_font(caption.add_run(f"圖 {entry.number}　{entry.title}"), bold=True)


# ── Table rendering ───────────────────────────────────────────────────────────

def render_table(doc: Document, tables: list[TableEntry], marker_idx: int) -> None:
    """Render the table referenced by a [TABLE:n] marker."""
    if marker_idx >= len(tables):
        return
    entry = tables[marker_idx]
    if not entry.rows:
        return

    col_count = max(len(row) for row in entry.rows)

    caption = doc.add_paragraph(style="TblCaption")
    _para_fmt(caption, align=WD_ALIGN_PARAGRAPH.CENTER, space_before_pt=6)
    _set_font(caption.add_run(f"表 {entry.number}　{entry.title}"), bold=True)

    tbl = doc.add_table(rows=len(entry.rows), cols=col_count)
    tbl.style = "Table Grid"
    for r_idx, row_data in enumerate(entry.rows):
        for c_idx, cell_text in enumerate(row_data):
            if c_idx >= col_count:
                break
            cell = tbl.rows[r_idx].cells[c_idx]
            cell.text = ""
            para = cell.paragraphs[0]
            _para_fmt(para, align=WD_ALIGN_PARAGRAPH.CENTER)
            _set_font(para.add_run(cell_text), size_pt=11)

    doc.add_paragraph()


# ── Section / chapter builders ────────────────────────────────────────────────

def build_section(doc: Document, sec: Section, depth: int,
                  tables: list[TableEntry] | None = None,
                  figures: list[FigureEntry] | None = None) -> None:
    label = f"{sec.id}　{sec.titleZh}"
    if depth == 1:
        _heading1(doc, label)
    elif depth == 2:
        _heading2(doc, label)
    else:
        para = doc.add_paragraph(style="Heading 3")
        _para_fmt(para, align=WD_ALIGN_PARAGRAPH.LEFT,
                  left_indent_pt=BODY_PT * 2, space_before_pt=4)
        _set_font(para.add_run(label), size_pt=BODY_PT, bold=True)

    for line in sec.content.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("[TABLE:") and stripped.endswith("]") and tables:
            try:
                render_table(doc, tables, int(stripped[7:-1]))
            except ValueError:
                _body(doc, line)
        elif stripped.startswith("[FIGURE:") and stripped.endswith("]") and figures:
            try:
                render_figure(doc, figures, int(stripped[8:-1]))
            except ValueError:
                _body(doc, line)
        else:
            _body(doc, line)

    if sec.subsections:
        for sub in sec.subsections:
            build_section(doc, sub, depth + 1, tables, figures)


def build_chapters(doc: Document, chapters: list[Chapter],
                   tables: list[TableEntry] | None = None,
                   figures: list[FigureEntry] | None = None) -> None:
    for ch in chapters:
        _add_page_break(doc)
        _chapter_heading(doc, f"{_to_zh(ch.number)}、　{ch.titleZh}", size_pt=16)
        if ch.titleEn:
            _center(doc, f"Chapter {_to_roman(ch.number)}: {ch.titleEn}",
                    size_pt=14, bold=True, italic=True, space_after=8)
        for sec in ch.sections:
            build_section(doc, sec, 1, tables, figures)


# ── Back matter builders ──────────────────────────────────────────────────────

def build_bibliography(doc: Document, bibliography: list[str], chapter_number: int) -> None:
    _add_page_break(doc)
    _chapter_heading(doc, f"{_to_zh(chapter_number)}、　參考文獻", size_pt=16)
    for ref in bibliography:
        if ref.strip():
            para = doc.add_paragraph()
            _para_fmt(para, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                      hanging_pt=BODY_PT * 2, space_after_pt=4)
            _set_font(para.add_run(ref))


def build_appendices(doc: Document, appendices: list[Appendix]) -> None:
    for ap in appendices:
        _add_page_break(doc)
        _center(doc, f"附錄{ap.label}　{ap.title}", size_pt=16, bold=True, space_after=8)
        for para_text in ap.content.split("\n"):
            if para_text.strip():
                _body(doc, para_text)
