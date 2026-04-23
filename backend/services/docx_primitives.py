"""Low-level python-docx utilities with no thesis-specific knowledge."""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt

FONT_ZH = "標楷體"
FONT_EN = "Times New Roman"
BODY_PT = 12


# ── Run / paragraph formatting ────────────────────────────────────────────────

def _set_font(run, latin=FONT_EN, east_asia=FONT_ZH, size_pt=BODY_PT,
              bold=False, italic=False) -> None:
    run.font.name = latin
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    rPr = run._r.get_or_add_rPr()
    existing = rPr.find(qn("w:rFonts"))
    if existing is None:
        existing = OxmlElement("w:rFonts")
        rPr.insert(0, existing)
    existing.set(qn("w:ascii"), latin)
    existing.set(qn("w:hAnsi"), latin)
    existing.set(qn("w:eastAsia"), east_asia)


def _para_fmt(para, align=WD_ALIGN_PARAGRAPH.LEFT, first_indent_pt=0,
              left_indent_pt=0, hanging_pt=0,
              space_before_pt=0, space_after_pt=0) -> None:
    fmt = para.paragraph_format
    fmt.alignment = align
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = 1.5
    if first_indent_pt:
        fmt.first_line_indent = Pt(first_indent_pt)
    if left_indent_pt:
        fmt.left_indent = Pt(left_indent_pt)
    if hanging_pt:
        fmt.first_line_indent = Pt(-hanging_pt)
        fmt.left_indent = Pt(hanging_pt)
    if space_before_pt:
        fmt.space_before = Pt(space_before_pt)
    if space_after_pt:
        fmt.space_after = Pt(space_after_pt)


# ── Page / section setup ──────────────────────────────────────────────────────

def _add_page_number_field(run) -> None:
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    run._r.append(fldChar1)

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "
    run._r.append(instrText)

    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar2)


def _set_section_page_number(section, fmt="lowerRoman", start=1) -> None:
    sectPr = section._sectPr
    for existing in sectPr.findall(qn("w:pgNumType")):
        sectPr.remove(existing)
    pgNumType = OxmlElement("w:pgNumType")
    pgNumType.set(qn("w:fmt"), fmt)
    pgNumType.set(qn("w:start"), str(start))
    sectPr.append(pgNumType)


def _setup_section(section) -> None:
    """A4 size with 1-inch margins."""
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)


def _add_footer_page_number(section) -> None:
    section.footer.is_linked_to_previous = False
    para = section.footer.paragraphs[0]
    para.clear()
    _para_fmt(para, align=WD_ALIGN_PARAGRAPH.CENTER)
    run = para.add_run()
    _set_font(run, size_pt=10)
    _add_page_number_field(run)


def _add_page_break(doc: Document) -> None:
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(WD_BREAK.PAGE)


# ── Paragraph factories ───────────────────────────────────────────────────────

def _center(doc: Document, text: str, size_pt=BODY_PT, bold=False, italic=False,
            space_before=0, space_after=0, style: str | None = None) -> None:
    para = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    _para_fmt(para, align=WD_ALIGN_PARAGRAPH.CENTER,
              space_before_pt=space_before, space_after_pt=space_after)
    run = para.add_run(text)
    _set_font(run, size_pt=size_pt, bold=bold, italic=italic)


def _body(doc: Document, text: str, indent=True) -> None:
    para = doc.add_paragraph()
    _para_fmt(para, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
              first_indent_pt=(BODY_PT * 2) if indent else 0)
    run = para.add_run(text)
    _set_font(run)


def _heading1(doc: Document, text: str) -> None:
    para = doc.add_paragraph(style="Heading 1")
    _para_fmt(para, align=WD_ALIGN_PARAGRAPH.LEFT, space_before_pt=8)
    run = para.add_run(text)
    _set_font(run, size_pt=13, bold=True)


def _heading2(doc: Document, text: str) -> None:
    para = doc.add_paragraph(style="Heading 2")
    _para_fmt(para, align=WD_ALIGN_PARAGRAPH.LEFT,
              left_indent_pt=BODY_PT, space_before_pt=4)
    run = para.add_run(text)
    _set_font(run, size_pt=BODY_PT)


def _chapter_heading(doc: Document, text: str, size_pt: int, italic: bool = False) -> None:
    """Heading 1 style but centred, so Word's TOC picks it up."""
    para = doc.add_paragraph(style="Heading 1")
    _para_fmt(para, align=WD_ALIGN_PARAGRAPH.LEFT, space_after_pt=6)
    run = para.add_run(text)
    _set_font(run, size_pt=size_pt, bold=True, italic=italic)


# ── Number conversion utilities ───────────────────────────────────────────────

_ZH_NUMS = ["零", "一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]


def _to_zh(n: int) -> str:
    if n <= 10:
        return _ZH_NUMS[n]
    if n < 20:
        return f"十{_ZH_NUMS[n - 10]}"
    return str(n)


_ROMAN_VALS = [1000, 900, 500, 400, 100, 90, 50, 40, 10, 9, 5, 4, 1]
_ROMAN_SYMS = ["M", "CM", "D", "CD", "C", "XC", "L", "XL", "X", "IX", "V", "IV", "I"]


def _to_roman(n: int) -> str:
    result = ""
    for v, s in zip(_ROMAN_VALS, _ROMAN_SYMS):
        while n >= v:
            result += s
            n -= v
    return result
