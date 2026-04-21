"""Orchestrates all builders to produce the final NCU thesis .docx."""
import io

from docx import Document
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from schemas.thesis_schema import ThesisSchema
from services.body_builder import (
    build_appendices, build_bibliography, build_chapters,
)
from services.docx_primitives import (
    FONT_EN, FONT_ZH, _add_footer_page_number, _add_page_break,
    _set_section_page_number, _setup_section,
)
from services.front_matter_builder import (
    build_abstract_en, build_abstract_zh, build_acknowledgments,
    build_cover, build_figure_list, build_placeholder, build_symbols,
    build_table_list, build_toc,
)


# ── Word document style setup (NCU-specific, called once per document) ────────

def _configure_heading_styles(doc: Document) -> None:
    specs = [("Heading 1", 16, True), ("Heading 2", 13, True), ("Heading 3", 12, True)]
    for name, size_pt, bold in specs:
        try:
            style = doc.styles[name]
        except KeyError:
            continue
        style.font.name = FONT_EN
        style.font.size = Pt(size_pt)
        style.font.bold = bold
        style.font.color.rgb = RGBColor(0, 0, 0)
        rPr = style.element.get_or_add_rPr()
        for existing in rPr.findall(qn("w:rFonts")):
            rPr.remove(existing)
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), FONT_EN)
        rFonts.set(qn("w:hAnsi"), FONT_EN)
        rFonts.set(qn("w:eastAsia"), FONT_ZH)
        rPr.insert(0, rFonts)


def _configure_toc_styles(doc: Document) -> None:
    for level in (1, 2, 3):
        try:
            style = doc.styles[f"TOC {level}"]
        except KeyError:
            continue
        style.font.name = FONT_EN
        style.font.size = Pt(12)
        style.font.color.rgb = RGBColor(0, 0, 0)
        rPr = style.element.get_or_add_rPr()
        for existing in rPr.findall(qn("w:rFonts")):
            rPr.remove(existing)
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), FONT_EN)
        rFonts.set(qn("w:hAnsi"), FONT_EN)
        rFonts.set(qn("w:eastAsia"), FONT_ZH)
        rPr.insert(0, rFonts)


def _enable_update_fields_on_open(doc: Document) -> None:
    settings = doc.settings.element
    for existing in settings.findall(qn("w:updateFields")):
        settings.remove(existing)
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    settings.append(update)


# ── Public API ────────────────────────────────────────────────────────────────

def build_thesis_docx(schema: ThesisSchema) -> bytes:
    doc = Document()
    _configure_heading_styles(doc)
    _configure_toc_styles(doc)
    _enable_update_fields_on_open(doc)

    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    # ── Section 1: front matter (roman page numbers) ──────────────────────────
    sec1 = doc.sections[0]
    _setup_section(sec1)
    _set_section_page_number(sec1, fmt="lowerRoman", start=1)
    _add_footer_page_number(sec1)

    build_cover(doc, schema)
    _add_page_break(doc)
    build_cover(doc, schema)           # duplicate cover required by NCU
    build_placeholder(doc, "授權書")
    build_placeholder(doc, "指導教授推薦書")
    build_placeholder(doc, "口試委員審定書")
    build_abstract_zh(doc, schema)
    build_abstract_en(doc, schema)
    build_acknowledgments(doc, schema)
    build_toc(doc)
    build_figure_list(doc)
    build_table_list(doc)
    if schema.symbols:
        build_symbols(doc, schema.symbols)

    # ── Section 2: body (arabic page numbers, restart from 1) ─────────────────
    sec2 = doc.add_section(WD_SECTION.NEW_PAGE)
    _setup_section(sec2)
    _set_section_page_number(sec2, fmt="decimal", start=1)
    _add_footer_page_number(sec2)

    build_chapters(doc, schema.chapters, schema.tables)
    build_bibliography(doc, schema.bibliography)
    if schema.appendices:
        build_appendices(doc, schema.appendices)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
