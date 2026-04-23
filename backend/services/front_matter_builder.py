"""Front matter section builders: cover, abstracts, acknowledgments, TOC, symbols."""
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from schemas.thesis_schema import SymbolEntry, ThesisSchema
from services.docx_primitives import (
    BODY_PT, _add_page_break, _body, _center, _para_fmt, _set_font,
)


# ── TOC field helper (used only within this module) ───────────────────────────

def _insert_toc_field(doc: Document, switches: str) -> None:
    """Insert a Word TOC field that updates on first open (F9)."""
    para = doc.add_paragraph()
    _para_fmt(para)
    run = para.add_run()
    _set_font(run)

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" TOC {switches} "
    run._r.append(instr)

    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    run._r.append(sep)

    placeholder = OxmlElement("w:t")
    placeholder.text = "（請於 Word 中按 F9 更新目錄）"
    run._r.append(placeholder)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(end)


# ── Section builders ──────────────────────────────────────────────────────────

def build_cover(doc: Document, schema: ThesisSchema) -> None:
    c = schema.cover
    degree_zh = f"{c.degree}論文"
    degree_en = "Doctoral Dissertation" if c.degree == "博士" else "Master's Thesis"

    for _ in range(2):
        doc.add_paragraph()
    _center(doc, "國立中央大學", size_pt=18, bold=True)
    _center(doc, c.department, size_pt=16, bold=True)
    doc.add_paragraph()
    _center(doc, degree_zh, size_pt=16, bold=True)
    _center(doc, degree_en, size_pt=14, bold=True)
    for _ in range(2):
        doc.add_paragraph()
    _center(doc, c.titleZh, size_pt=18, bold=True)
    _center(doc, c.titleEn, size_pt=14, bold=True, italic=True)
    for _ in range(3):
        doc.add_paragraph()
    _center(doc, f"研 究 生：{c.studentName}", size_pt=14)
    _center(doc, f"指導教授：{c.advisorName}", size_pt=14)
    for _ in range(3):
        doc.add_paragraph()
    _center(doc, f"中華民國 {c.year} 年 {c.month} 月", size_pt=14)


def build_placeholder(doc: Document, title: str) -> None:
    _add_page_break(doc)
    for _ in range(8):
        doc.add_paragraph()
    _center(doc, f"【{title}】", size_pt=16, bold=True)
    doc.add_paragraph()
    _center(doc, "（請貼附正式文件）", size_pt=12)


def build_abstract_zh(doc: Document, schema: ThesisSchema) -> None:
    _add_page_break(doc)
    _center(doc, "摘要", size_pt=16, bold=True, space_after=8, style="FMTitle")
    _body(doc, schema.abstractZh.content)
    doc.add_paragraph()
    para = doc.add_paragraph()
    _para_fmt(para, first_indent_pt=BODY_PT * 2)
    r1 = para.add_run("關鍵詞：")
    _set_font(r1, bold=True)
    r2 = para.add_run("、".join(schema.abstractZh.keywords))
    _set_font(r2)


def build_abstract_en(doc: Document, schema: ThesisSchema) -> None:
    _add_page_break(doc)
    _center(doc, schema.cover.titleEn, size_pt=14, bold=True, italic=True, space_after=4)
    _center(doc, "Abstract", size_pt=14, bold=True, space_after=8, style="FMTitle")
    _body(doc, schema.abstractEn.content)
    doc.add_paragraph()
    para = doc.add_paragraph()
    _para_fmt(para, first_indent_pt=BODY_PT * 2)
    r1 = para.add_run("Keywords: ")
    _set_font(r1, bold=True)
    r2 = para.add_run(", ".join(schema.abstractEn.keywords))
    _set_font(r2)


def build_acknowledgments(doc: Document, schema: ThesisSchema) -> None:
    _add_page_break(doc)
    _center(doc, "誌謝", size_pt=16, bold=True, space_after=8, style="FMTitle")
    for para_text in schema.acknowledgments.split("\n"):
        if para_text.strip():
            _body(doc, para_text)


def build_toc(doc: Document) -> None:
    _add_page_break(doc)
    _center(doc, "目錄", size_pt=16, bold=True, space_after=8, style="FMTitle")
    _insert_toc_field(doc, r'\t "FMTitle,1" \o "1-3" \h \z')


def build_figure_list(doc: Document) -> None:
    _add_page_break(doc)
    _center(doc, "圖目錄", size_pt=16, bold=True, space_after=8, style="FMTitle")
    _insert_toc_field(doc, r'\t "FigCaption,1" \h \z')


def build_table_list(doc: Document) -> None:
    _add_page_break(doc)
    _center(doc, "表目錄", size_pt=16, bold=True, space_after=8, style="FMTitle")
    _insert_toc_field(doc, r'\t "TblCaption,1" \h \z')


def build_symbols(doc: Document, symbols: list[SymbolEntry]) -> None:
    _add_page_break(doc)
    _center(doc, "符號說明", size_pt=16, bold=True, space_after=8, style="FMTitle")
    for s in symbols:
        para = doc.add_paragraph()
        _para_fmt(para)
        r1 = para.add_run(f"{s.symbol}\t")
        _set_font(r1)
        r2 = para.add_run(s.description)
        _set_font(r2)
